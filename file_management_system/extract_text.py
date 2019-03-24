#!/usr/bin/python
# -*- coding: utf-8 -*-

from pdfminer.converter import PDFPageAggregator
from pdfminer.pdfparser import PDFParser
from pdfminer.pdfdocument import PDFDocument
from pdfminer.pdfpage import PDFPage
from pdfminer.pdfpage import PDFTextExtractionNotAllowed
from pdfminer.pdfinterp import PDFResourceManager
from pdfminer.pdfinterp import PDFPageInterpreter
from pdfminer.layout import *
import re
import docx
import openpyxl
import xlrd
import numpy as np


def readPdf(pdfName):
    # 打开一个pdf文件
    fp = open(pdfName, 'rb')
    # 创建一个PDF文档解析器对象
    parser = PDFParser(fp)
    # 创建一个PDF文档对象存储文档结构
    # 提供密码初始化，没有就不用传该参数
    # document = PDFDocument(parser, password)
    document = PDFDocument(parser)
    # 检查文件是否允许文本提取
    if not document.is_extractable:
        raise PDFTextExtractionNotAllowed
    # 创建一个PDF资源管理器对象来存储共享资源
    # caching = False不缓存
    rsrcmgr = PDFResourceManager(caching=False)
    # 创建一个PDF设备对象
    laparams = LAParams()
    # 创建一个PDF页面聚合对象
    device = PDFPageAggregator(rsrcmgr, laparams=laparams)
    # 创建一个PDF解析器对象
    interpreter = PDFPageInterpreter(rsrcmgr, device)
    # 处理文档当中的每个页面

    fullText = []
    replace = re.compile(r'\s+');
    # 循环遍历列表，每次处理一个page的内容
    for page in PDFPage.create_pages(document):
        interpreter.process_page(page)
        # 接受该页面的LTPage对象
        layout = device.get_result()
        # 这里layout是一个LTPage对象 里面存放着 这个page解析出的各种对象
        # 一般包括LTTextBox, LTFigure, LTImage, LTTextBoxHorizontal 等等
        for x in layout:
            # 如果x是水平文本对象的话
            if (isinstance(x, LTTextBoxHorizontal)):
                text = re.sub(replace, '', x.get_text())
                if len(text) != 0:
                    fullText.append(text)
    return '\n'.join(fullText)


def readWord(docName):
    fullText = []
    doc = docx.Document(docName)
    paras = doc.paragraphs
    for p in paras:
        fullText.append(p.text)
    return '\n'.join(fullText)


"""
def readExcel(xlsName):
    fullText = ""
    wb = openpyxl.load_workbook(xlsName)
    sheet = wb[wb.sheetnames[0]]
    for row in sheet.iter_rows(min_row=sheet.min_row, max_col=sheet.max_row):
        for cell in row:
            if cell.value is not None:
                fullText = fullText + " " + str(cell.value)
    return fullText
"""


def readExcel(xlsName):
    wb = xlrd.open_workbook(xlsName)
    sheet = wb.sheet_by_index(0)
    nrows = sheet.nrows
    a = []
    for row in range(nrows):
        a.extend(sheet.row_values(row))
    fullText = " ".join(np.array(a).astype(str))
    return fullText
